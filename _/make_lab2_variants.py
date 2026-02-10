from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

from make_labs_common_functions import replace_skip, parse_variants_file, lab2_variants_redef

def convert_to_math_unicode_formula(text):
    text = text.replace('_{1}', '₁')#.replace('_1', '₁')
    text = text.replace('_{2}', '₂')#.replace('_2', '₂')
    text = text.replace('_{3}', '₃')#.replace('_3', '₃')
    text = text.replace('^{2}', '²')#.replace('^2', '²')
    text = text.replace('^{3}', '³')#.replace('^3', '³')
    text = text.replace('^{4}', '⁴')#.replace('^4', '⁴')
    text = text.replace('_{i}', 'ᵢ')#.replace('_i', 'ᵢ')
    text = text.replace('_{j}', 'ⱼ')#.replace('_j', 'ⱼ')
    text = text.replace('_{ij}', 'ᵢⱼ')#.replace('_ij', 'ᵢⱼ')
    text = text.replace('_{2ij}', '₂ᵢⱼ')#.replace('_2ij', '₂ᵢⱼ')
    text = text.replace("^{'}", "'")
    
    return text

def make_lab2_variants_and_save_in_docx(variants, year, group, output_file):
    doc = Document()
    
    table = doc.add_table(rows=len(variants)*2, cols=4)
    table.style = 'Table Grid'
    
    for i, variant in enumerate(variants):
        row_top = table.rows[i*2]
        row_bottom = table.rows[i*2+1]
        table.cell(i*2, 0).merge(table.cell(i*2 + 1, 0)) # Об'єднання сусідніх комірок першого стовпця
   
        cell_num = row_top.cells[0]
        cell_num.text = 'Варіант №' + str(variant['number']) + '\nКІ-' + str(group) + ', ' + str(year - 1) + '/' + str(year) + ' н.р.'
        cell_num.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        cell_formula = row_top.cells[1]
        cell_formula.merge(row_top.cells[2])
        cell_formula.merge(row_top.cells[3])
        cell_formula.text = convert_to_math_unicode_formula(variant['formula']) + ", " + variant['type']
        cell_formula.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        cell_bi = row_bottom.cells[1]
        cell_bi.text = convert_to_math_unicode_formula(variant['b_i'].replace('; ', ',\n').replace('; ',',\n'))
        cell_bi.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        cell_a1 = row_bottom.cells[2]
        cell_a1.text = 'y₂=' + convert_to_math_unicode_formula(variant['y2'])
        cell_a1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        cell_a2c = row_bottom.cells[3]
        cell_a2c.text = 'Y₃=' + convert_to_math_unicode_formula(variant['Y3']) + '\n' + convert_to_math_unicode_formula(variant['C2_ij'])
        cell_a2c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
    
    doc.save(output_file)
    print("Created:", output_file)

if __name__ == "__main__":
    YEAR_FIRST = 2026
    YEAR_LAST = 2027
    GROUP_FIRST = 301
    GROUP_LAST = 309
    LAB2_VARIANTS_DATA_FILE_NAME = "lab2_variants_data.txt"

    # Читаємо базові варіанти
    variants_ = parse_variants_file(LAB2_VARIANTS_DATA_FILE_NAME)

    for year in range(YEAR_FIRST, YEAR_LAST + 1):
        for group in range(GROUP_FIRST, GROUP_LAST + 1):
            # Створюємо нові варіанти
            variants = lab2_variants_redef(variants_, year, group)
            
            # Створюємо шлях до файлу
            filename = "PRO_LAB2_VARIANTSANDMATLAB/" + str(year - 1) + "_" + str(year) + "/KI" + str(group) + "/l2_variants_ki" + str(group) + "_" + str(year - 1) + str(year) + ".docx"
            path = Path(filename)            
            path.parent.mkdir(parents=True, exist_ok=True) # if path.parent != Path("."):

            make_lab2_variants_and_save_in_docx(
                variants,
                year,
                group, 
                filename                         
            )
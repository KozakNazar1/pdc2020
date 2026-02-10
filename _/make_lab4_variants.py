import random
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from pathlib import Path

from make_labs_common_functions import get_permutation_count, get_permutation_by_number

def make_lab4_variants_and_save_in_docx(year, group, variant_count, output_file):
    n, k = 10, 9
    permutation_count = get_permutation_count(n, k)
    years_block = 12;
    group_max_count = 14;
    permutation_part_area = permutation_count // (variant_count + 1) // years_block // group_max_count

    # Створення документа
    doc = Document()
    
    # Додавання заголовку
    title = doc.add_heading(f'ВАРІАНТИ ЗАВДАНЬ (КІ-{group}, {year - 1}/{year} н.р.)\n', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Створення таблиці: (кількість варіантів + заголовк) x 2 рядки кожен
    table = doc.add_table(rows=2*(variant_count + 1), cols=4)
    table.style = 'Table Grid'
    
    # Заповнення заголовків таблиці
    # Перший рядок
    table.cell(0, 0).text = "Варіанти №"
    table.cell(0, 0).merge(table.cell(1, 0)) # Об'єднання сусідніх комірок першого стовпця у заголовку
    table.cell(0, 1).text = "Граф що відображає конфігурацію зв'язків між процесами"
    table.cell(0, 1).merge(table.cell(0, 3)) # Інше об'єднання комірок для заголовку
    
    # Другий рядок - підзаголовки
    table.cell(1, 1).text = "N1"
    table.cell(1, 2).text = "N2"
    table.cell(1, 3).text = "N3"
    
    # Заповнення даних для кожного варіанту
    for index in range(0, variant_count):
        row_offset = 2 + index * 2
        
        # Номер варіанту
        table.cell(row_offset, 0).text = str(index + 1)
        table.cell(row_offset, 0).merge(table.cell(row_offset + 1, 0)) # Об'єднання сусідніх комірок першого стовпця
        
        # Граф
        permutation_index = permutation_part_area * (index + 1) * ((year - 2025) %years_block) * ((group - 299)%group_max_count)
        #permutation_index %= permutation_count # !
        permutation = get_permutation_by_number(permutation_index, 10, [0, 1, 2, 3, 4, 5, 6, 7, 8, 9])
        table.cell(row_offset, 1).text = f" → {permutation[0]} → {permutation[1]} → {permutation[2]} → {permutation[3]} → {permutation[4]} → {permutation[5]} → {permutation[6]} → {permutation[7]} → {permutation[8]} →\n└─────────────────────────┘"
        table.cell(row_offset, 1).merge(table.cell(row_offset, 3))
        
        # Рядок з числами N1, N2, N3
        numbers_row = row_offset + 1
        for j in range(3):
            table.cell(numbers_row, j + 1).text = str(permutation_index * permutation_index * (j + 1) % 999)
    
    # Налаштування стилів таблиці
    for row in table.rows:
        for cell in row.cells:
            # Центрування тексту
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Жирний шрифт для заголовків
            if row == table.rows[0] or row == table.rows[1]:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(11)
    
    # Налаштування ширини стовпців
    table.columns[0].width = Inches(0.8)    # Вариант№
    table.columns[1].width = Inches(2.5)    # Граф і N1
    table.columns[2].width = Inches(0.8)    # N2
    table.columns[3].width = Inches(0.8)    # N3
  
    # Збереження документу
    doc.save(output_file)
    print("Created:", output_file)

if __name__ == "__main__":
    YEAR_FIRST = 2026
    YEAR_LAST = 2027
    GROUP_FIRST = 301
    GROUP_LAST = 309
    VARIANT_COUNT = 30

    for year in range(YEAR_FIRST, YEAR_LAST + 1):
        for group in range(GROUP_FIRST, GROUP_LAST + 1):            
            # Створюємо шлях до файлу
            filename = "PRO_LAB4_VARIANTS/" + str(year - 1) + "_" + str(year) + "/KI" + str(group) + "/l4_variants_ki" + str(group) + "_" + str(year - 1) + str(year) + ".docx"
            path = Path(filename)            
            path.parent.mkdir(parents=True, exist_ok=True) # if path.parent != Path("."):

            make_lab4_variants_and_save_in_docx(
                year,
                group,
                VARIANT_COUNT,
                filename                         
            )
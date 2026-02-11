from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import os
from pathlib import Path

def set_cell_width(cell, width):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width.cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def make_lab3_variants_and_save_in_docx(
    year,
    group,
    variant_count,
    img_dir,
    output_file
):
    doc = Document()

    # Заголовок
    title = doc.add_paragraph("ВАРІАНТИ ЗАВДАНЬ (КІ-" + str(group) + ", " + str(year - 1) + "/" + str(year) + " н.р.)")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True

    doc.add_paragraph(
        "Матриця A задається однозначно і залежить лише від розмірності даних.\n"
        "Для матриці B: заштрихована область – довільні цілі числа, відмінні від нуля, "
        "а незаштрихована область – нулі."
    )

    # Таблиця
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Заголовки
    hdr = table.rows[0].cells
    hdr[0].text = "Варіант №"
    hdr[1].text = "Тип матриці A"
    hdr[2].text = "Тип матриці B"

    for cell in hdr:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True

    # Фіксовані ширини колонок
    widths = [Cm(2), Cm(7), Cm(7)]

    # Заповнення рядків
    for i in range(1, variant_count + 1):
        row = table.add_row().cells

        row[0].text = str(i)
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        param_Y = year - 2023;
        param_G = group - 295;

        imgA = os.path.join(img_dir, f"{(i - 1 + param_Y + param_G // 7) % variant_count + 1}_1.jpg")
        imgB = os.path.join(img_dir, f"{(i - 1 + param_G) % variant_count + 1}_2.jpg")

        if os.path.exists(imgA):
            p = row[1].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(imgA, width=Cm(1.8))

        if os.path.exists(imgB):
            p = row[2].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(imgB, width=Cm(1.8))

        for cell, w in zip(row, widths):
            set_cell_width(cell, w)

    # Ширини для заголовка
    for cell, w in zip(hdr, widths):
        set_cell_width(cell, w)

    doc.save(output_file)
    print("Created:", output_file)

if __name__ == "__main__":
    YEAR = 2026

    IMG_DIR = "image_l3"
    YEAR_FIRST = 2026
    YEAR_LAST = 2027
    GROUP_FIRST = 301
    GROUP_LAST = 309
    VARIANT_COUNT = 30

    for year in range(YEAR_FIRST, YEAR_LAST + 1):
        for group in range(GROUP_FIRST, GROUP_LAST + 1):
            # Створюємо шлях до файлу
            filename = "PRO_LAB3_VARIANTS/" + str(year - 1) + "_" + str(year) + "/KI" + str(group) + "/l3_variants_ki" + str(group) + "_" + str(year - 1) + str(year) + ".docx"
            path = Path(filename)            
            path.parent.mkdir(parents=True, exist_ok=True) # if path.parent != Path("."):

            make_lab3_variants_and_save_in_docx(
                year=year,
                group=group,
                variant_count=VARIANT_COUNT,
                img_dir=IMG_DIR,
                output_file=filename
            )
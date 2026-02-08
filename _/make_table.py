from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_to_unicode(text): # convert_to_normal_formula
    text = text.replace('_{1}', '₁')#.replace('_1', '₁')
    text = text.replace('_{2}', '₂')#.replace('_2', '₂')
    text = text.replace('_{3}', '₃')#.replace('_3', '₃')
    text = text.replace('^{2}', '²')#.replace('^2', '²')
    text = text.replace('^{3}', '³')#.replace('^3', '³')
    text = text.replace('^{4}', '⁴')#.replace('^4', '⁴')
    text = text.replace('_{i}', 'ᵢ')#.replace('_i', 'ᵢ')
    text = text.replace('_{j}', 'ⱼ')#.replace('_j', 'ⱼ')
    text = text.replace('_{ij}', 'ᵢⱼ')#.replace('_ij', 'ᵢⱼ')
    text = text.replace('_{2ij}', '₂ᵢⱼ')#.replace('_2ij', '₂ᵢⱼ')
    text = text.replace("^{'}", "'")
    
    return text

def read_variants_file(filename):
    variants = []
    with open(filename, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith('#'):
                continue
            
            parts = line.split('|')
            if len(parts) == 7:
                number = int(parts[0])
                formula = parts[1]
                type_text = parts[2]
                b_i = parts[3].replace('; ', '\n\n')
                y2 = parts[4]
                Y3 = parts[5]
                C2_ij = parts[6]
                
                variants.append({
                    'number': number,
                    'formula': formula + ', ' + type_text,
                    'b_i': b_i,
                    'y2': 'y₂=' + y2,
                    'Y3': 'Y₃=' + Y3 + '\n\n' + C2_ij
                })
    return variants

def create_doc_from_variants(year, group, variants, output_file):
    doc = Document()
    
    table = doc.add_table(rows=len(variants)*2, cols=4)
    table.style = 'Table Grid'
    
    for i, variant in enumerate(variants):
        row_top = table.rows[i*2]
        
        cell_num = row_top.cells[0]
        cell_num.text = 'Варіант №' + str(variant['number'])
        cell_num.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        cell_formula = row_top.cells[1]
        cell_formula.merge(row_top.cells[2])
        cell_formula.merge(row_top.cells[3])
        cell_formula.text = convert_to_unicode(variant['formula'])
        cell_formula.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        row_bottom = table.rows[i*2+1]
        
        cell_empty = row_bottom.cells[0]
        cell_empty.text = group + ', ' + str(int(year) - 1) + '/' + year + ' н.р.'
        
        cell_bi = row_bottom.cells[1]
        cell_bi.text = convert_to_unicode(variant['b_i'])
        cell_bi.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        cell_a1 = row_bottom.cells[2]
        cell_a1.text = convert_to_unicode(variant['y2'])
        cell_a1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        cell_a2c = row_bottom.cells[3]
        cell_a2c.text = convert_to_unicode(variant['Y3'])
        cell_a2c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
    
    doc.save(output_file)
    print("Document created:", output_file)

if __name__ == "__main__":
    variants = read_variants_file("variants_data.txt")
    
    #with open("variants_processed.py", "w", encoding='utf-8') as f:
    #    f.write("variants = [\n")
    #    for i, v in enumerate(variants):
    #        f.write(f'    {{\n')
    #        f.write(f'        "number": {v["number"]},\n')
    #        f.write(f'        "formula": "{v["formula"]}",\n')
    #        f.write(f'        "b_i": "{v["b_i"]}",\n')
    #        f.write(f'        "a1": "{v["a1"]}",\n')
    #        f.write(f'        "Y3": "{v["Y3"]}"\n')
    #        f.write(f'    }}')
    #        if i < len(variants) - 1:
    #            f.write(",")
    #        f.write("\n")
    #    f.write("]\n")
    #print("Processed data saved: variants_processed.py")
    
    year = '2026'
    group = 'КІ-308'
    create_doc_from_variants(year, group, variants, "V_LAB1__.docx")